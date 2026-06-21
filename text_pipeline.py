"""
Citrinitas · 熔知 — 文本管线

OCR / 文本提取 / 分块 / 嵌入。
所有函数返回统一的 {"ok": bool, ...} 格式。
"""

import os
import re
import io
import json
import hashlib
import subprocess
import tempfile
import logging

import requests
from docx import Document
from bs4 import BeautifulSoup

from qconst import PROJECT_DIR, IMAGES_DIR, EMBED_MODEL, OLLAMA_URL

logger = logging.getLogger(__name__)

# ── PIL 可选 ──
try:
    from PIL import Image as PILImage
    HAS_PIL = True
except ImportError:
    PILImage = None
    HAS_PIL = False

# Tesseract 备选（可通过环境变量覆盖）
_TESSERACT_FALLBACK = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
TESSERACT = os.environ.get("KB_TESSERACT_PATH") or _TESSERACT_FALLBACK
os.environ.setdefault("TESSDATA_PREFIX", os.environ.get("KB_TESSDATA_PREFIX") or r"D:\Tesseract-OCR\tessdata")

# PaddleOCR 主力引擎（延迟初始化）
_paddle_ocr = None

# PPStructureV3 引擎（延迟初始化）
_structure_engine = None


# ═══════════════════════════════════════════
# OCR 引擎初始化
# ═══════════════════════════════════════════

def _get_paddle():
    """延迟初始化 PaddleOCR（首次调用才加载模型）"""
    global _paddle_ocr
    if _paddle_ocr is None:
        try:
            from paddleocr import PaddleOCR
            _paddle_ocr = PaddleOCR(lang='ch', ocr_version='PP-OCRv4', use_textline_orientation=True)
        except ImportError:
            raise ImportError(
                "PaddleOCR 未安装。运行: "
                "D:/uv/tools/private-gpt/Scripts/python.exe -m pip install paddlepaddle paddleocr"
            )
    return _paddle_ocr


def _get_structure_engine():
    """延迟初始化 PPStructureV3（首次调用才加载模型）"""
    global _structure_engine
    if _structure_engine is None:
        try:
            from paddleocr import PPStructureV3
            _structure_engine = PPStructureV3(
                lang='ch',
                use_formula_recognition=True,
                use_table_recognition=True,
                use_chart_recognition=False,
                format_block_content=True,
            )
        except ImportError as e:
            raise ImportError(
                f"PPStructureV3 初始化失败: {e}\n"
                f"请确保 paddlex[ocr] 已安装：\n"
                f"  D:/uv/tools/private-gpt/Scripts/python.exe -m pip install 'paddlex[ocr]==3.7.0'"
            )
    return _structure_engine


# ═══════════════════════════════════════════
# OCR 实现
# ═══════════════════════════════════════════

def _ocr_paddle(image_path: str) -> dict:
    """
    PaddleOCR 识别（主力引擎，中文+公式优化）
    返回: {"ok": true, "text": "...", "chars": N, "conf": 0.95, "raw": [...]}
    """
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"图片不存在: {image_path}")

    engine = _get_paddle()
    result = engine.predict(image_path)

    if not result or not isinstance(result, list) or not result[0]:
        return {"ok": True, "text": "", "chars": 0, "conf": 0.0, "raw": []}

    page = result[0]
    texts = page.get('rec_texts', [])
    scores = page.get('rec_scores', [])

    lines = []
    confs = []
    for i, text in enumerate(texts):
        text = text.strip()
        if text:
            lines.append(text)
            conf = scores[i] if i < len(scores) else 0.0
            confs.append(float(conf))

    full_text = "\n".join(lines)
    avg_conf = sum(confs) / len(confs) if confs else 0.0

    return {
        "ok": True,
        "text": full_text,
        "chars": len(full_text),
        "conf": round(float(avg_conf), 4),
        "lines": len(lines),
        "raw": [{"text": t, "conf": c} for t, c in zip(lines, confs)]
    }


def _ocr_tesseract(image_path: str, lang: str = "chi_sim+eng") -> dict:
    """
    Tesseract OCR（备选引擎）
    返回: {"ok": true, "text": "...", "chars": N, "conf": null}
    """
    if not os.path.exists(TESSERACT):
        raise FileNotFoundError(f"Tesseract 未找到: {TESSERACT}")
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"图片不存在: {image_path}")

    outbase = image_path + "_ocr_tmp"
    result = subprocess.run(
        [TESSERACT, image_path, outbase, "-l", lang, "--psm", "6"],
        capture_output=True, text=True, timeout=120
    )
    if result.returncode != 0:
        raise RuntimeError(f"Tesseract 失败: {result.stderr.strip()}")

    txt_path = outbase + ".txt"
    with open(txt_path, "r", encoding="utf-8") as f:
        text = f.read().strip()
    os.remove(txt_path)

    return {
        "ok": True,
        "text": text,
        "chars": len(text),
        "conf": None,
        "lines": text.count("\n") + 1,
        "raw": []
    }


def _html_table_to_markdown(html_text: str) -> str:
    """将 HTML <table> 转换为 Markdown 表格（简化版，处理 PPStructureV3 输出）"""

    def _convert_table(match):
        table_html = match.group(0)
        rows = re.findall(r'<tr>(.*?)</tr>', table_html, re.DOTALL)
        md_rows = []
        for row_idx, row in enumerate(rows):
            cells = re.findall(r'<(?:td|th).*?>(.*?)</(?:td|th)>', row, re.DOTALL | re.IGNORECASE)
            clean_cells = []
            for c in cells:
                c = re.sub(r'<img\s+src="([^"]+)"[^>]*>', r'[image: \1]', c)
                c = re.sub(r'<div[^>]*>', ' ', c)
                c = re.sub(r'</div>', '', c)
                c = re.sub(r'<[^>]+>', '', c)
                c = c.strip()
                c = c.replace('|', '\\|')
                clean_cells.append(c)
            md_rows.append('| ' + ' | '.join(clean_cells) + ' |')
            if row_idx == 0:
                md_rows.append('|' + '|'.join([' --- ' for _ in clean_cells]) + '|')
        return '\n'.join(md_rows)

    html_text = re.sub(r'(?:<div[^>]*>\s*)?<html><body>\s*<table[^>]*>.*?</table>\s*</body></html>\s*(?:</div>)?',
                        _convert_table, html_text, flags=re.DOTALL)
    return html_text


def _assemble_blocks_v2(blocks: list, page_idx: int, image_list: list, source_image: str, page) -> str:
    """
    手动拼装 PPStructureV3 LayoutBlocks（_to_markdown 失败时的回退路径）。
    blocks 是 LayoutBlock 对象列表（label/content/bbox）。
    """
    if not blocks:
        return ""

    parts = []
    for b in blocks:
        label = getattr(b, 'label', 'text')
        content = getattr(b, 'content', '') or ''
        content = str(content)

        if label == 'table':
            parts.append(f"\n[表格] {content}\n")
        elif label in ('figure', 'figure_title', 'image'):
            imgs = page.get('imgs_in_doc', [])
            if imgs and hasattr(b, 'bbox'):
                bbox = b.bbox
                for img_item in imgs:
                    if isinstance(img_item, dict) and 'coordinate' in img_item:
                        coord = img_item['coordinate']
                        if (abs(bbox[0] - coord[0]) < 30 and abs(bbox[1] - coord[1]) < 30):
                            pil = img_item.get('img')
                            if pil:
                                _ensure_images_dir()
                                dest_name = f"fig_p{page_idx}_{len(image_list)}.png"
                                dest_path = os.path.join(IMAGES_DIR, dest_name)
                                pil.save(dest_path)
                                image_list.append(dest_path)
                                parts.append(f"[image: {dest_path}]")
                            break
            if not parts or parts[-1].startswith('[image') is False:
                parts.append(f"[图表] {content}")
        elif label == 'formula':
            content = content.strip()
            if content and not content.startswith('$$'):
                content = f"$${content}$$"
            parts.append(content)
        else:
            if content.strip():
                parts.append(content.strip())

    return "\n".join(parts)


def _check_ocr_quality(ocr_result: dict, image_path: str = None) -> dict:
    """
    检查 OCR 识别质量。

    检测信号:
    - 中文占比 < 10% → 可能图片模糊/纯英文/非文档
    - 字符数过少 → 可能图片模糊
    - 重复乱码 → 图片质量差
    - 平均置信度低 → 不确定

    返回: {"grade": "good|warn|bad", "score": 0-100, "issues": [...], "suggestion": "..."}
    """
    text = ocr_result.get("text", "")
    conf = ocr_result.get("conf")
    issues = []

    # 1. 中文占比
    cjk_chars = len(re.findall(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]', text))
    total_chars = len(text.strip())
    cjk_ratio = cjk_chars / max(total_chars, 1)

    if total_chars < 3:
        issues.append("识别文字过少（<3字符）")
    elif total_chars < 20:
        issues.append(f"识别文字较少（{total_chars}字符），可能图片模糊")

    if cjk_ratio < 0.05 and total_chars > 10:
        issues.append(f"中文占比极低（{cjk_ratio:.1%}），确认是否中文文档")

    # 2. 乱码检测（连续非打印字符）
    garbled = re.findall(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]{3,}', text)
    if garbled:
        issues.append("检测到乱码字符")

    # 3. 置信度
    if conf is not None:
        if conf < 0.5:
            issues.append(f"平均置信度低（{conf:.2f}），OCR 不确定")
        elif conf < 0.7:
            issues.append(f"平均置信度偏低（{conf:.2f}）")

    # 评分
    if not issues:
        score = 90
        grade = "good"
    elif len(issues) == 1 and ("偏低" in issues[0] or "较少" in issues[0]):
        score = 65
        grade = "warn"
    else:
        score = 30
        grade = "bad"

    return {
        "grade": grade,
        "score": score,
        "chars": total_chars,
        "cjk_ratio": round(cjk_ratio, 3),
        "issues": issues,
        "suggestion": "可直接入库" if grade == "good"
                       else "建议人工审核后入库" if grade == "warn"
                       else "图片可能模糊或非中文文档，建议重新拍摄"
    }


def _ocr_structured(image_path: str) -> dict:
    """
    PPStructureV3 结构化识别（主力引擎）
    自动检测版面区域：文字/公式/表格/图表，分别用最优方式处理。

    PPStructureV3.predict() 返回 list[LayoutParsingResultV2]。
    使用内置 _to_markdown() 获取统一 Markdown，包含：
      - HTML 表格 → 转为 Markdown 表格
      - 内嵌 LaTeX 公式 ($...$)
      - <img> 图片引用 → 保存为 [image: path]

    返回: {
        "ok": true,
        "text": "...",
        "chars": N,
        "blocks": [...],
        "images": [...],
        "conf": 0.95,
    }
    """
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"图片不存在: {image_path}")

    engine = _get_structure_engine()

    try:
        pages = engine.predict(image_path)
    except Exception as e:
        return {"ok": False, "error": f"PPStructureV3 识别失败: {e}"}

    if not pages:
        return {"ok": True, "text": "", "chars": 0, "blocks": [], "images": [], "conf": 0.0}

    all_text_parts = []
    all_images = []
    block_summary = []

    for page_idx, page in enumerate(pages):
        try:
            md_result = page._to_markdown(pretty=True, show_formula_number=False)
        except Exception:
            layout_blocks = page.get('parsing_res_list', [])
            page_text = _assemble_blocks_v2(layout_blocks, page_idx, all_images, image_path, page)
            all_text_parts.append(page_text)
            block_summary.append({"page": page_idx, "type": "fallback", "length": len(page_text)})
            continue

        md_text = md_result.get('markdown_texts', '')
        if not md_text:
            continue

        # ── 1. 收集并保存图片 ──
        path_to_pil = {}
        imgs_in_doc = page.get('imgs_in_doc', [])
        for img_item in imgs_in_doc:
            if isinstance(img_item, dict):
                p = img_item.get('path', '')
                pil = img_item.get('img')
                if p and pil:
                    path_to_pil[p] = pil

        def _replace_img(match):
            src = match.group(1)
            _ensure_images_dir()
            dest_name = f"fig_p{page_idx}_{len(all_images)}.png"
            dest_path = os.path.join(IMAGES_DIR, dest_name)
            if src in path_to_pil:
                path_to_pil[src].save(dest_path)
            all_images.append(dest_path)
            return f"\n[image: {dest_path}]\n"

        md_text = re.sub(
            r'(?:<div[^>]*>\s*)?<img\s+src="([^"]+)"[^>]*>(?:\s*</div>)?',
            _replace_img, md_text
        )

        # ── 2. 转换 HTML 表格为 Markdown 表格 ──
        md_text = _html_table_to_markdown(md_text)

        # ── 3. 清理残余 HTML 标签 ──
        md_text = re.sub(r'<div[^>]*>', '\n', md_text)
        md_text = re.sub(r'</div>', '', md_text)
        md_text = re.sub(r'<html><body>', '', md_text)
        md_text = re.sub(r'</body></html>', '', md_text)
        md_text = re.sub(r'\n{3,}', '\n\n', md_text)

        all_text_parts.append(md_text.strip())
        block_summary.append({"page": page_idx, "type": "structured", "length": len(md_text)})

    full_text = "\n\n".join(all_text_parts)

    return {
        "ok": True,
        "text": full_text,
        "chars": len(full_text),
        "blocks": block_summary,
        "images": all_images,
        "conf": 0.95,
    }


# ═══════════════════════════════════════════
# 公共 OCR 入口
# ═══════════════════════════════════════════

def ocr_image(image_path: str) -> dict:
    """
    公共 OCR 入口：自动选择最优引擎识别图片文字。

    优先使用 PPStructureV3（结构化识别：文字+表格+公式），
    回退到 PaddleOCR（基础文字识别）。

    返回:
        {
            "ok": true,
            "ocr_text": "识别出的文字...",
            "text": "识别出的文字...",
            "chars": N,
            "conf": 0.95,
            "needs_correction": false,
            "quality": {...}
        }
    """
    if not os.path.exists(image_path):
        return {"ok": False, "error": f"图片不存在: {image_path}"}

    # 优先尝试 PPStructureV3
    try:
        result = _ocr_structured(image_path)
        if result.get("ok") and result.get("text"):
            quality = _check_ocr_quality(result, image_path)
            return {
                "ok": True,
                "ocr_text": result.get("text", ""),
                "text": result.get("text", ""),
                "chars": result.get("chars", len(result.get("text", ""))),
                "conf": result.get("conf", 0.0),
                "needs_correction": quality.get("grade") != "good",
                "quality": quality,
                "model": "PPStructureV3",
            }
    except Exception as e:
        logger.warning(f"[OCR] PPStructure 失败，回退到 PaddleOCR: {e}")

    # 回退到 PaddleOCR
    try:
        result = _ocr_paddle(image_path)
        if result.get("ok"):
            quality = _check_ocr_quality(result, image_path)
            return {
                "ok": True,
                "ocr_text": result.get("text", ""),
                "text": result.get("text", ""),
                "chars": result.get("chars", len(result.get("text", ""))),
                "conf": result.get("conf", 0.0),
                "needs_correction": quality.get("grade") != "good",
                "quality": quality,
                "model": "PaddleOCR",
            }
            }
    except Exception as e:
        return {"ok": False, "error": f"OCR 引擎初始化失败: {e}"}

    return {"ok": False, "error": "无法加载任何 OCR 引擎"}


# ═══════════════════════════════════════════
# 文本提取
# ═══════════════════════════════════════════

def extract_text(file_path: str) -> dict:
    """
    统一文本提取入口：根据文件扩展名调用对应解析器。

    支持格式：
      - .txt / .md / .json / .csv → 直接读取（自动检测编码）
      - .docx → python-docx 提取文本
      - .html / .htm → BeautifulSoup 提取文本（去除标签）
      - .srt → 解析 SRT 字幕格式（去除时间戳）
      - .pdf → 尝试 pdfplumber 提取文本

    返回:
        {"ok": True, "text": "...", "chars": N, "meta": {}}
        {"ok": False, "error": "..."}
    """
    if not os.path.exists(file_path):
        return {"ok": False, "error": f"文件不存在: {file_path}"}

    ext = os.path.splitext(file_path)[1].lower()

    # ── 纯文本格式：直接读取 ──
    if ext in (".txt", ".md", ".json", ".csv", ".log"):
        encoding = detect_encoding(file_path)
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read()
            return {"ok": True, "text": text, "chars": len(text), "meta": {"encoding": encoding}}
        except Exception as e:
            return {"ok": False, "error": f"读取文件失败: {e}"}

    # ── DOCX 格式 ──
    if ext == ".docx":
        try:
            doc = Document(file_path)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            text = "\n\n".join(parts)
            return {"ok": True, "text": text, "chars": len(text), "meta": {"format": "docx"}}
        except Exception as e:
            return {"ok": False, "error": f"解析 DOCX 失败: {e}"}

    # ── HTML 格式 ──
    if ext in (".html", ".htm"):
        try:
            with open(file_path, "r", encoding=detect_encoding(file_path)) as f:
                soup = BeautifulSoup(f.read(), "html.parser")
            for tag in soup(["script", "style"]):
                tag.decompose()
            text = soup.get_text(separator="\n")
            text = re.sub(r"\n{3,}", "\n\n", text).strip()
            return {"ok": True, "text": text, "chars": len(text), "meta": {"format": "html"}}
        except Exception as e:
            return {"ok": False, "error": f"解析 HTML 失败: {e}"}

    # ── SRT 字幕格式 ──
    if ext == ".srt":
        try:
            encoding = detect_encoding(file_path)
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()
            lines = []
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.isdigit():
                    continue
                if "-->" in line:
                    continue
                lines.append(line)
            text = "\n".join(lines)
            return {"ok": True, "text": text, "chars": len(text), "meta": {"format": "srt", "encoding": encoding}}
        except Exception as e:
            return {"ok": False, "error": f"解析 SRT 失败: {e}"}

    # ── PDF 格式 ──
    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        parts.append(page_text)
                text = "\n\n".join(parts)
            if not text.strip():
                return {"ok": False, "error": "PDF 无文本内容（可能是扫描件，需要 OCR）"}
            return {"ok": True, "text": text, "chars": len(text), "meta": {"format": "pdf", "pages": len(pdf.pages)}}
        except ImportError:
            return {"ok": False, "error": "需要安装 pdfplumber：pip install pdfplumber"}
        except Exception as e:
            return {"ok": False, "error": f"解析 PDF 失败: {e}"}

    # ── 不支持的格式 ──
    return {"ok": False, "error": f"不支持的文件格式: {ext}"}


# ═══════════════════════════════════════════
# 编码 & 语言检测
# ═══════════════════════════════════════════

def detect_encoding(file_path: str, sample_size: int = 10000) -> str:
    """
    检测文件编码。优先 chardet，失败后用 UTF-8 → GBK → latin-1 兜底链。
    sample_size: 用于检测的字节数（默认 10000，约 10KB）
    """
    with open(file_path, "rb") as f:
        raw = f.read(sample_size)
    if not raw:
        return "utf-8"  # 空文件，默认 UTF-8

    # 先试 chardet（如果已安装）
    try:
        import chardet
        result = chardet.detect(raw)
        enc = result.get("encoding", "").strip().lower()
        conf = result.get("confidence", 0)
        if enc and conf >= 0.6:
            enc_map = {
                "utf-8": "utf-8",
                "ascii": "utf-8",
                "gb2312": "gbk",
                "gbk": "gbk",
                "gb18030": "gb18030",
                "big5": "big5",
                "iso-8859-1": "latin-1",
                "windows-1252": "cp1252",
            }
            return enc_map.get(enc, enc)
    except ImportError:
        pass  # chardet 未安装，走兜底链

    # 兜底链：UTF-8 → GBK → latin-1
    try:
        raw.decode("utf-8")
        return "utf-8"
    except UnicodeDecodeError:
        pass
    try:
        raw.decode("gbk")
        return "gbk"
    except UnicodeDecodeError:
        pass
    # 最后兜底：latin-1 永不失败（但可能乱码）
    return "latin-1"


def _detect_language(text: str) -> str:
    """通过 Unicode 区块统计检测语言（前 2000 字）。"""
    sample = text[:2000]
    if not sample.strip():
        return "zh"
    total = len(sample)
    cjk = sum(1 for c in sample if '\u4e00' <= c <= '\u9fff' or '\u3400' <= c <= '\u4dbf')
    hiragana = sum(1 for c in sample if '\u3040' <= c <= '\u309f')
    katakana = sum(1 for c in sample if '\u30a0' <= c <= '\u30ff')
    hangul = sum(1 for c in sample if '\uac00' <= c <= '\ud7af')
    latin = sum(1 for c in sample if c.isascii() and c.isalpha())

    cjk_ratio = cjk / total
    ja_ratio = (hiragana + katakana) / total
    ko_ratio = hangul / total
    en_ratio = latin / total

    if cjk_ratio >= 0.30:
        return "zh"
    if ja_ratio >= 0.10:
        return "ja"
    if ko_ratio >= 0.10:
        return "ko"
    if en_ratio >= 0.60:
        return "en"
    return "zh"  # 兜底


def detect_language(text: str) -> str:
    """
    程序检测文本语言（中/英），不调用 LLM，确定性输出。

    逻辑：
        - 统计 CJK 统一汉字范围（\u4e00-\u9fff）字符占比
        - 占比 > 30% → "zh"
        - 否则 → "en"（默认英文）
        - 空文本 → "en"

    返回:
        "zh" | "en" | "ja" | "ko"  (远期可扩展)
    """
    if not text:
        return "en"
    total = len(text)
    if total == 0:
        return "en"
    cjk_count = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    if cjk_count / total > 0.3:
        return "zh"
    return "en"


# ═══════════════════════════════════════════
# 嵌入
# ═══════════════════════════════════════════

def _embed(texts: list[str], model: str = EMBED_MODEL) -> list[list[float]]:
    """调用 Ollama 批量获取嵌入向量（优先批量，失败回退逐条）。"""
    if not texts:
        return []
    # 尝试批量 API（Ollama /api/embed 支持 input 数组）
    try:
        resp = requests.post(
            f"{OLLAMA_URL}/api/embed",
            json={"model": model, "input": texts},
            timeout=120
        )
        resp.raise_for_status()
        embeddings = resp.json().get("embeddings", [])
        if embeddings and len(embeddings) == len(texts):
            return embeddings
    except Exception as e:
        logger.warning(f"[Embed] 批量嵌入失败，回退到逐条: {e}")
    # 逐条回退
    vectors = []
    for i, text in enumerate(texts):
        try:
            resp = requests.post(
                f"{OLLAMA_URL}/api/embeddings",
                json={"model": model, "prompt": text},
                timeout=60
            )
            resp.raise_for_status()
            vectors.append(resp.json()["embedding"])
        except Exception:
            if len(texts) == 1:
                raise
            logger.warning(f"[Embed] 块 #{i} 嵌入失败（{len(text)} 字符），已跳过")
    return vectors


# ═══════════════════════════════════════════
# 文本切块
# ═══════════════════════════════════════════

def _text_hash(text: str) -> str:
    """内容的去重哈希（规范化后 SHA256，取前 32 位）"""
    normalized = re.sub(r'\s+', ' ', text).strip().lower()
    return hashlib.sha256(normalized.encode()).hexdigest()[:32]


def _extract_images(text: str) -> list[str]:
    """提取文本中的图片引用（支持 3 种格式：[:image:] / Markdown / HTML）。"""
    images = []
    images.extend(re.findall(r'\[image:\s*([^\]]+)\]', text))
    images.extend(re.findall(r'!\[.*?\]\(([^\)]+)\)', text))
    images.extend(re.findall(r'<img[^>]+src=["\']([^"\']+)["\']', text, re.IGNORECASE))
    seen = set()
    unique = []
    for img in images:
        if img not in seen:
            seen.add(img)
            unique.append(img)
    return unique


def _ensure_images_dir():
    """确保图片存储目录存在"""
    os.makedirs(IMAGES_DIR, exist_ok=True)


def _safe_slice_point(text: str, target: int) -> int:
    """
    寻找安全的切片点：优先在 target 附近找标点/空格，避免切断中文。
    向前搜索 50 字符，向后搜索 50 字符。
    找不到则返回 target（允许切断）。
    """
    if target <= 0 or target >= len(text):
        return target
    start = max(0, target - 50)
    end = min(len(text), target + 50)
    punctuation = set('。！？；;,.!? \n\t')
    for i in range(target, start, -1):
        if text[i] in punctuation:
            return i + 1
    for i in range(target, end):
        if text[i] in punctuation:
            return i + 1
    for i in range(target, start, -1):
        if text[i] == ' ':
            return i + 1
    for i in range(target, end):
        if text[i] == ' ':
            return i + 1
    return target


def _split_long_paragraph(text: str, max_chars: int, overlap: int) -> list[str]:
    """将长段落按句子切分，不切断内联公式 $...$。
    超长句（>max_chars）安全切片，避免切断中文。"""
    sentences = re.split(r'(?<=[。；;])\s*', text)
    chunks = []
    current = ""
    for sent in sentences:
        sent = sent.strip()
        if not sent:
            continue
        if len(current) + len(sent) < max_chars:
            current = (current + sent).strip()
        else:
            if current:
                chunks.append(current)
            if len(sent) > max_chars:
                while len(sent) > max_chars:
                    cut = _safe_slice_point(sent, max_chars)
                    chunks.append(sent[:cut])
                    sent = sent[cut:].strip()
                current = sent if sent else ""
            else:
                current = sent
    if current:
        chunks.append(current)
    if overlap > 0 and len(chunks) > 1:
        overlapped = [chunks[0]]
        for i in range(1, len(chunks)):
            prev_tail = chunks[i-1][-overlap:] if len(chunks[i-1]) >= overlap else chunks[i-1]
            overlapped.append(prev_tail + "\n\n" + chunks[i])
        return overlapped
    return chunks


def _chunk_text(text: str, max_chars: int = 800, overlap: int = 60) -> list[str]:
    """
    将文本切分为重叠的块。
    保护原子结构（公式/表格/图片引用）不被截断。
    """
    # ── 第1步：保护原子块（替换为占位符）──
    placeholders = {}
    counter = [0]

    def _protect(match):
        key = f"__ATOMIC_{counter[0]}__"
        placeholders[key] = match.group(0)
        counter[0] += 1
        return key

    text = re.sub(r'\$\$[\s\S]*?\$\$', _protect, text)
    text = re.sub(r'(?:^\|.+\|$\n?)+', _protect, text, flags=re.MULTILINE)
    text = re.sub(r'\[image:[^\]]+\]', _protect, text)
    text = re.sub(r'\[图表\][^\n]*', _protect, text)

    # ── 第2步：正常切分 ──
    paragraphs = re.split(r'\n\s*\n', text)
    chunks = []
    current = ""
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current) + len(para) < max_chars:
            current = (current + "\n\n" + para).strip()
        else:
            if current:
                chunks.append(current)
            if len(para) > max_chars:
                sub_chunks = _split_long_paragraph(para, max_chars, overlap)
                chunks.extend(sub_chunks)
                current = ""
            else:
                current = para
    if current:
        chunks.append(current)

    # ── 第3步：还原占位符 ──
    restored = []
    for chunk in chunks:
        for key, val in placeholders.items():
            chunk = chunk.replace(key, val)
        restored.append(chunk)

    return restored
